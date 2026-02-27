#!/usr/bin/env python3
import sys
sys.path.insert(0, '/Users/palvashamadireddy/Downloads/CPP2_project/CPP2/underwater_segmentation/venv/lib/python3.11/site-packages')

from docx import Document

def analyze_full_docx(filepath):
    doc = Document(filepath)
    print(f"\n=== Full Analysis: {filepath} ===")
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    
    # Count total characters
    total_chars = 0
    for para in doc.paragraphs:
        total_chars += len(para.text)
    
    print(f"Total characters: {total_chars}")
    
    # Print ALL paragraphs structure
    print("\n--- Full structure (non-empty paragraphs only) ---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else "Normal"
            # Get runs to check font info
            runs_info = ""
            if para.runs:
                run = para.runs[0]
                font_name = run.font.name if run.font.name else "No font"
                runs_info = f" | Font: {font_name}"
            print(f"{i}: [{style}]{runs_info}")
            # Print first 200 chars of each paragraph
            print(f"   {text[:200]}")

# Analyze both V3 files
analyze_full_docx('/Users/palvashamadireddy/Downloads/CPP2_project/CPP2/underwater_segmentation/RESEARCH_PAPER_FINAL_V3.docx')
analyze_full_docx('/Users/palvashamadireddy/Downloads/CPP2_project/CPP2/underwater_segmentation/PROJECT_REPORT_FINAL_V3.docx')
