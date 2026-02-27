#!/usr/bin/env python3
import sys
sys.path.insert(0, '/Users/palvashamadireddy/Downloads/CPP2_project/CPP2/underwater_segmentation/venv/lib/python3.14/site-packages')

from docx import Document

def analyze_docx(filepath):
    doc = Document(filepath)
    print(f"\n=== Analyzing: {filepath} ===")
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    
    # Count total characters
    total_chars = 0
    for para in doc.paragraphs:
        total_chars += len(para.text)
    
    print(f"Total characters: {total_chars}")
    
    # Print first 50 paragraphs to understand structure
    print("\n--- First 50 paragraphs structure ---")
    for i, para in enumerate(doc.paragraphs[:50]):
        text = para.text.strip()
        if text:
            # Get style name
            style = para.style.name if para.style else "Normal"
            # Get runs to check font info
            runs_info = ""
            if para.runs:
                run = para.runs[0]
                font_name = run.font.name if run.font.name else "No font"
                font_size = str(run.font.size) if run.font.size else "No size"
                runs_info = f" | Font: {font_name}, Size: {font_size}"
            print(f"{i}: [{style}]{runs_info}")
            # Print first 100 chars of each paragraph
            print(f"   {text[:150]}...")

# Analyze both V3 files
analyze_docx('/Users/palvashamadireddy/Downloads/CPP2_project/CPP2/underwater_segmentation/RESEARCH_PAPER_FINAL_V3.docx')
analyze_docx('/Users/palvashamadireddy/Downloads/CPP2_project/CPP2/underwater_segmentation/PROJECT_REPORT_FINAL_V3.docx')
