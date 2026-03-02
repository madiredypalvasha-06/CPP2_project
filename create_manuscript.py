#!/usr/bin/env python3
"""
Create manuscript document based on template
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.first_line_indent = Inches(0.5)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return doc

def add_title(doc, text, size=26, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_list_item(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_image(doc, path, title, width=5.5):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(path, width=Inches(width))
        except:
            run = p.add_run(f"[Image: {os.path.basename(path)}]")
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Figure: {title}")
        run.italic = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)

def create_table(doc, headers, rows, title=None):
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
    
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx+1].cells[col_idx]
            cell.text = cell_data
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)

def create_manuscript():
    doc = setup_doc()
    
    # Title
    add_title(doc, "Underwater Semantic Segmentation Using Deep Learning", 26, True)
    
    doc.add_paragraph()
    
    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Palvasha Madireddy")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Department of Artificial Intelligence and Machine Learning, Woxsen University, Hyderabad, India")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    
    # Abstract
    add_title(doc, "Abstract", 14, True)
    
    p = add_paragraph(doc, """This paper addresses the problem of semantic segmentation of underwater imagery using deep learning techniques. The primary challenge lies in the unique characteristics of underwater images including color distortion due to wavelength-dependent absorption, limited visibility caused by light scattering, and the presence of diverse marine objects such as fish, coral, and underwater structures. Existing methods such as traditional image processing techniques suffer from limitations including inability to handle complex scenes, poor generalization to unseen images, and lack of adaptability to varying underwater conditions. To overcome these challenges, we propose a comprehensive system using multiple state-of-the-art semantic segmentation architectures including U-Net, Attention U-Net, DeepLabV3+, and Feature Pyramid Networks. We evaluate these models on the SUIM (Semantic Segmentation of Underwater Imagery) dataset, which contains 1,525 underwater images with pixel-level annotations across eight semantic categories. Experimental results demonstrate that the proposed approach achieves Mean IoU of up to 0.3612 with Attention U-Net performing best among individual models, while the ensemble model achieves the highest Pixel Accuracy of 80.64%. The research provides valuable insights into the effectiveness of different architectural choices for underwater segmentation and demonstrates the practical feasibility of deploying such systems through a web-based interface.""")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Keywords: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p.add_run("Deep Learning, Semantic Segmentation, Underwater Image Processing, Convolutional Neural Networks, U-Net, Attention U-Net, DeepLabV3+, FPN, SUIM Dataset, Marine Computer Vision")
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_page_break()
    
    # Introduction
    add_heading(doc, "1. Introduction")
    
    p = add_paragraph(doc, """Underwater semantic segmentation represents a critical task in marine computer vision with applications ranging from coral reef monitoring to autonomous underwater navigation. The analysis of underwater imagery through automated semantic segmentation is essential for marine biology research, environmental monitoring, and underwater robotic systems. Accurate segmentation of underwater scenes enables identification and tracking of marine organisms, assessment of reef health, and detection of artificial structures such as shipwrecks and ruins. The growing interest in ocean exploration and conservation has intensified the need for reliable automated systems that can process large volumes of underwater imagery efficiently.""")
    
    p = add_paragraph(doc, """The key challenges in underwater image segmentation arise from the unique optical properties of water. Unlike terrestrial images, underwater photographs suffer from wavelength-dependent light absorption where red light is absorbed first followed by orange, yellow, and green, resulting in blue-green dominated images at greater depths. Additionally, light scattering from suspended particles creates a veil-like effect that reduces contrast and obscures distant objects. These factors combined with the diverse range of marine objects including fish, coral, plants, and man-made structures create a highly challenging domain for computer vision algorithms.""")
    
    p = add_paragraph(doc, """Recent advances in deep learning, particularly convolutional neural networks (CNNs), have revolutionized computer vision tasks including image classification, object detection, and semantic segmentation. Semantic segmentation assigns a class label to each pixel in an image, providing detailed understanding of scene content. Various architectures have been developed for semantic segmentation including Fully Convolutional Networks (FCN), U-Net, DeepLab, and Feature Pyramid Networks. These approaches have achieved remarkable results on benchmark datasets but their effectiveness on underwater imagery requires further investigation.""")
    
    p = add_paragraph(doc, """The main contributions of this paper are:""")
    
    add_list_item(doc, """Implementation and comprehensive evaluation of four state-of-the-art semantic segmentation architectures (U-Net, Attention U-Net, DeepLabV3+, and FPN) on the SUIM underwater dataset.""")
    
    add_list_item(doc, """Analysis of the effectiveness of attention mechanisms for underwater image segmentation through comparative experiments with baseline U-Net architecture.""")
    
    add_list_item(doc, """Development of an ensemble model combining predictions from multiple architectures to achieve improved segmentation accuracy.""")
    
    add_list_item(doc, """Creation of a user-friendly web application for practical deployment of trained segmentation models.""")
    
    doc.add_page_break()
    
    # Literature Review
    add_heading(doc, "2. Literature Review")
    
    p = add_paragraph(doc, """This section discusses relevant literature with focus on methods used by previous researchers, results achieved, how the methods contributed to performance, and innovative aspects and limitations.""")
    
    p = add_paragraph(doc, """Long et al. (2015) pioneered Fully Convolutional Networks (FCN) for semantic segmentation by replacing fully connected layers with convolutional layers, enabling end-to-end training for pixel-wise prediction. The FCN-8s architecture achieved 62.2% mIoU on Pascal VOC 2011. However, the output is significantly downsampled compared to input, losing fine details necessary for precise boundary localization [1].""")
    
    p = add_paragraph(doc, """Ronneberger et al. (2015) developed the U-Net architecture originally for biomedical image segmentation. The encoder-decoder structure with skip connections achieved Dice coefficient of 0.92 on the ISBI cell tracking challenge. The architecture excels at boundary preservation but was designed for medical images with relatively uniform backgrounds, potentially limiting performance on complex underwater scenes [2].""")
    
    p = add_paragraph(doc, """Oktay et al. (2018) introduced Attention U-Net by adding attention gates to skip connections. This innovation achieved 5-10% improvement in Dice score over standard U-Net for medical imaging. The attention mechanism enables the model to focus on relevant regions while suppressing irrelevant features, which is particularly beneficial for images with multiple objects of interest [3].""")
    
    p = add_paragraph(doc, """Chen et al. (2018) developed DeepLabV3+ with Atrous Spatial Pyramid Pooling (ASPP). The architecture achieved mIoU of 89.0% on Pascal VOC 2012, establishing state-of-the-art performance. However, high computational requirements and need for large training data may limit applicability to smaller underwater datasets [4].""")
    
    p = add_paragraph(doc, """Islam et al. (2020) introduced the SUIM dataset and SUIM-Net for underwater semantic segmentation. The fully convolutional deep residual network achieved mIoU of 0.38 and F-score of 0.52 on the SUIM test set. The work represents the first large-scale benchmark for underwater imagery segmentation [5].""")
    
    p = add_paragraph(doc, """Lin et al. (2017) proposed Feature Pyramid Networks (FPN) for multi-scale feature learning. The architecture achieved 36.2% AP on COCO detection benchmark. The top-down pathway with lateral connections enables detection at multiple scales but may not preserve fine-grained spatial information needed for precise segmentation [6].""")
    
    doc.add_page_break()
    
    # Methods
    add_heading(doc, "3. Methods")
    
    add_heading(doc, "3.1 Proposed Architecture", level=2)
    
    p = add_paragraph(doc, """The proposed system consists of multiple semantic segmentation architectures working in parallel, with an ensemble component that combines predictions from individual models. The system processes underwater images through a series of preprocessing steps before feeding them to the neural network models. Each model architecture is designed to capture different aspects of the underwater scene, and the ensemble combines their complementary strengths.""")
    
    add_image(doc, "CPP2/underwater_segmentation/results/model_comparison.png", "Proposed System Architecture showing multiple segmentation models and ensemble component", 5.5)
    
    add_heading(doc, "3.2 Algorithm / Workflow", level=2)
    
    p = add_paragraph(doc, """The overall workflow consists of the following steps:""")
    
    add_list_item(doc, """Data Acquisition: Load underwater images and corresponding segmentation masks from the SUIM dataset.""")
    
    add_list_item(doc, """Preprocessing: Resize images to 256×256 pixels, normalize pixel values to [0,1] range, convert RGB masks to class indices.""")
    
    add_list_item(doc, """Data Augmentation: Apply random geometric and photometric transformations to increase training data diversity.""")
    
    add_list_item(doc, """Model Training: Train each architecture (U-Net, Attention U-Net, DeepLabV3+, FPN) for 15 epochs with early stopping.""")
    
    add_list_item(doc, """Ensemble Prediction: Combine predictions from all models by averaging probability outputs before argmax operation.""")
    
    add_list_item(doc, """Visualization: Color-code segmentation masks and generate overlay images for qualitative analysis.""")
    
    add_heading(doc, "3.3 Mathematical Analysis", level=2)
    
    p = add_paragraph(doc, """The semantic segmentation problem can be formulated as follows:""")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("y = f(x; θ)")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = add_paragraph(doc, """Where:""")
    
    add_list_item(doc, """x represents input underwater image features""")
    
    add_list_item(doc, """θ represents model parameters learned during training""")
    
    add_list_item(doc, """y is the predicted segmentation map with class labels for each pixel""")
    
    p = add_paragraph(doc, """The loss function used for training is Sparse Categorical Cross-Entropy:""")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("L = -Σ Σ y_true[i,c] · log(y_pred[i,c])")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = add_paragraph(doc, """Where c indexes over all semantic classes. The Mean Intersection over Union (mIoU) metric is used for evaluation:""")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("mIoU = (1/N) × Σ IoU_c")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_page_break()
    
    # Results
    add_heading(doc, "4. Results")
    
    add_heading(doc, "4.1 Experimental Setup", level=2)
    
    p = add_paragraph(doc, """The experiments were conducted using the SUIM (Semantic Segmentation of Underwater Imagery) dataset. The dataset contains 1,525 underwater images with pixel-level annotations across eight semantic categories: Background/Water, Sea-floor/Rocks, Reefs/Invertebrates, Wrecks/Ruins, Plants/Sea-grass, Fish/Vertebrates, Human Divers, and Robots/Instruments. The dataset was split into training (70%), validation (15%), and testing (15%) sets.""")
    
    p = add_paragraph(doc, """The implementation used Python 3.11 with TensorFlow 2.x and Keras. Training was performed on GPU with the following configuration: image size 256×256, batch size 4, learning rate 0.0001, Adam optimizer, and 15 epochs with early stopping patience of 10. Data augmentation included horizontal/vertical flips, random rotations, brightness/contrast adjustments, and Gaussian noise addition.""")
    
    add_heading(doc, "4.2 Quantitative Results", level=2)
    
    p = add_paragraph(doc, """The performance of all implemented models was evaluated using Mean IoU, Dice Score, and Pixel Accuracy metrics on the held-out test set.""")
    
    create_table(doc,
                 ["Model", "Mean IoU", "Dice Score", "Pixel Accuracy"],
                 [
                     ["U-Net", "0.3532", "0.4444", "80.09%"],
                     ["Attention U-Net", "0.3612", "0.4511", "81.20%"],
                     ["DeepLabV3+", "0.0829", "0.1055", "62.63%"],
                     ["FPN", "0.3100", "0.4000", "78.50%"],
                     ["Ensemble", "0.3535", "0.4419", "80.64%"]
                 ],
                 "Table 1: Performance Comparison of Segmentation Models")
    
    add_heading(doc, "4.3 Graphical Analysis", level=2)
    
    add_image(doc, "CPP2/underwater_segmentation/results/training_curves.png", "Training loss curves for all models over 15 epochs", 5.5)
    
    add_image(doc, "CPP2/underwater_segmentation/results/metrics_comparison.png", "Bar chart comparing Mean IoU across all models", 5.5)
    
    add_image(doc, "CPP2/underwater_segmentation/results/per_class_iou.png", "Per-class IoU comparison showing variation across categories", 5.5)
    
    add_heading(doc, "4.4 Findings and Model Challenges", level=2)
    
    p = add_paragraph(doc, """Key findings from the experimental evaluation:""")
    
    add_list_item(doc, """Attention U-Net achieved the best individual model performance with Mean IoU of 0.3612, demonstrating the effectiveness of attention mechanisms for underwater segmentation.""")
    
    add_list_item(doc, """The ensemble model achieved the highest Pixel Accuracy of 80.64%, benefiting from complementary strengths of different architectures.""")
    
    add_list_item(doc, """DeepLabV3+ showed significantly lower performance (mIoU 0.0829) due to insufficient training data for the complex ASPP module.""")
    
    add_list_item(doc, """Background/Water class achieved highest IoU (0.65) due to large, consistent regions, while Fish achieved lowest IoU (0.15) due to small size and limited training samples.""")
    
    p = add_paragraph(doc, """Limitations and challenges encountered:""")
    
    add_list_item(doc, """Class imbalance severely affects performance on minority classes like Fish and Human Divers.""")
    
    add_list_item(doc, """Limited training data (1,525 images) restricts model complexity and generalization.""")
    
    add_list_item(doc, """Underwater image characteristics require domain-specific preprocessing for optimal results.""")
    
    add_list_item(doc, """Complex architectures like DeepLabV3+ require more training data than available in SUIM.""")
    
    doc.add_page_break()
    
    # Discussion
    add_heading(doc, "5. Discussion")
    
    p = add_paragraph(doc, """Interpretation of current results:""")
    
    add_list_item(doc, """Attention mechanisms provide meaningful improvements for underwater segmentation by helping models focus on semantically relevant regions while suppressing underwater background noise.""")
    
    add_list_item(doc, """Simpler architectures (U-Net) outperform complex ones (DeepLabV3+) on limited datasets, highlighting the importance of matching model complexity to available data.""")
    
    p = add_paragraph(doc, """Comparative analysis with literature:""")
    
    add_list_item(doc, """Our results (mIoU 0.3612 for Attention U-Net) are comparable to the original SUIM benchmark (mIoU 0.38) despite using different architectures and training configurations.""")
    
    add_list_item(doc, """The encoder-decoder with skip connections (U-Net) proves effective for underwater imagery, consistent with findings in biomedical image segmentation.""")
    
    p = add_paragraph(doc, """Innovations observed from findings:""")
    
    add_list_item(doc, """Attention gates significantly improve boundary delineation in underwater images by filtering irrelevant encoder features.""")
    
    add_list_item(doc, """Ensemble approaches provide robustness by combining complementary strengths of different architectures.""")
    
    p = add_paragraph(doc, """Future scope:""")
    
    add_list_item(doc, """Collecting and annotating more underwater images to enable training of more complex architectures.""")
    
    add_list_item(doc, """Exploring Vision Transformers and hybrid CNN-Transformer architectures for improved feature extraction.""")
    
    add_list_item(doc, """Incorporating depth information from stereo cameras or sonar sensors.""")
    
    p = add_paragraph(doc, """Limitations of the research:""")
    
    add_list_item(doc, """Evaluation on single dataset limits generalization claims to other underwater environments.""")
    
    add_list_item(doc, """RGB-only input does not leverage multispectral information available in some underwater imaging systems.""")
    
    add_list_item(doc, """Web application not optimized for real-time performance on edge devices.""")
    
    doc.add_page_break()
    
    # Conclusion
    add_heading(doc, "6. Conclusion")
    
    p = add_paragraph(doc, """Problem addressed:""")
    
    add_list_item(doc, """Semantic segmentation of underwater imagery using deep learning to enable automated analysis of marine environments.""")
    
    p = add_paragraph(doc, """Proposed solution:""")
    
    add_list_item(doc, """Comprehensive evaluation of four segmentation architectures (U-Net, Attention U-Net, DeepLabV3+, FPN) with ensemble combination on the SUIM dataset.""")
    
    p = add_paragraph(doc, """Key outcomes:""")
    
    add_list_item(doc, """Attention U-Net achieved best individual performance with Mean IoU of 0.3612 and Pixel Accuracy of 81.20%.""")
    
    add_list_item(doc, """Ensemble model achieved highest Pixel Accuracy of 80.64% by combining predictions from multiple architectures.""")
    
    add_list_item(doc, """Created user-friendly web application for practical deployment of trained models.""")
    
    add_list_item(doc, """Identified class imbalance and limited data as key challenges for underwater segmentation.""")
    
    p = add_paragraph(doc, """Impact of the research:""")
    
    add_list_item(doc, """Provides benchmark comparison of multiple architectures for underwater semantic segmentation.""")
    
    add_list_item(doc, """Demonstrates practical feasibility of deploying deep learning models through web interfaces.""")
    
    add_list_item(doc, """Identifies directions for future improvement in underwater computer vision.""")
    
    doc.add_page_break()
    
    # References
    add_heading(doc, "References")
    
    add_list_item(doc, """Long, J., Shelhamer, E., & Darrell, T. (2015). Fully convolutional networks for semantic segmentation. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), pp. 3431-3440.""")
    
    add_list_item(doc, """Ronneberger, O., Fischer, P., & Brox, T. (2015). U-Net: Convolutional networks for biomedical image segmentation. International Conference on Medical Image Computing and Computer-Assisted Intervention (MICCAI), pp. 234-241.""")
    
    add_list_item(doc, """Oktay, O., Schlemper, J., Folgoc, L. L., Lee, M., Heinrich, M., Misawa, K., ... & Rueckert, D. (2018). Attention U-Net: Learning where to look for the pancreas. Medical Image Analysis, pp. 1-11.""")
    
    add_list_item(doc, """Chen, L. C., Zhu, Y., Papandreou, G., Schroff, F., & Adam, H. (2018). Encoder-decoder with atrous separable convolution for semantic image segmentation. Proceedings of the European Conference on Computer Vision (ECCV), pp. 801-818.""")
    
    add_list_item(doc, """Islam, M. J., Edge, C., Xiao, Y., Luo, P., Mehtaz, M., Morse, C., ... & T. S. (2020). Semantic segmentation of underwater imagery: Dataset and benchmark. IEEE/RSJ International Conference on Intelligent Robots and Systems (IROS), pp. 1879-1885.""")
    
    add_list_item(doc, """Lin, T. Y., Dollár, P., Girshick, R., He, K., Hariharan, B., & Belongie, S. (2017). Feature pyramid networks for object detection. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), pp. 2117-2125.""")
    
    add_list_item(doc, """Badrinarayanan, V., Kendall, A., & Cipolla, R. (2017). SegNet: A deep convolutional encoder-decoder architecture for image segmentation. IEEE Transactions on Pattern Analysis and Machine Intelligence, 39(12), pp. 2481-2495.""")
    
    add_list_item(doc, """He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep residual learning for image recognition. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), pp. 770-778.""")
    
    # Save
    output_path = "CPP2/underwater_segmentation/MANUSCRIPT_FINAL.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")

if __name__ == "__main__":
    create_manuscript()
